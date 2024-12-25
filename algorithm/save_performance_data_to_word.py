from docx import Document
import os

def save_performance_data_to_word(performance_data_list, flag= "", file_name='performance_data.docx'):


    document = Document()
    document.add_heading('Performance Data', level=1)

    # Create table with 3 columns
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Algorithm'
    hdr_cells[1].text = 'Time Taken (s)'
    hdr_cells[2].text = 'Memory Usage (MB)'

    # Add data rows
    for data in performance_data_list:
        row_cells = table.add_row().cells
        if flag == "":
            file_name = row_cells[0].text = data['function']
        row_cells[0].text = data['function']
        row_cells[1].text = f"{data['time_taken']:.6f}"
        row_cells[2].text = f"{data['memory_usage']:.6f}"
        
    if flag != "":
        file_name = flag
    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_name = file_name+".docx"
    file_path = os.path.join(current_dir, file_name)

    document.save(file_path)

# Example usage
if __name__ == "__main__":
    performance_data_list = [
        {'function': 'python_sort', 'time_taken': 0.19879969999965397, 'memory_usage': 7.9872},
        {'function': 'main_q_sort', 'time_taken': 0.19879969999965397, 'memory_usage': 7.9872},
        {'function': 'buble_sort', 'time_taken': 0.19879969999965397, 'memory_usage': 7.9872},
        {'function': 'q_sort', 'time_taken': 0.19879969999965397, 'memory_usage': 7.9872},
    ]
    save_performance_data_to_word(performance_data_list, flag="all")